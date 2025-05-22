# resume_app/views.py

import io, os, tempfile, json, re
from django.shortcuts import render, redirect
from django.http import HttpResponse, FileResponse
from django.conf import settings
from django.template.loader import render_to_string

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

from pdf2docx import Converter
from PyPDF2 import PdfReader
from openai import OpenAI, APIConnectionError
from weasyprint import HTML
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np

from .forms import ResumeForm, CoverLetterForm

# ————— Globals —————
client = OpenAI(api_key=settings.OPENAI_API_KEY)
BLACKLIST = {"json", "[]", ""}

# ————— Helper: TF-IDF fallback —————
def extract_keywords_tfidf(text, top_n=3):
    vec = TfidfVectorizer(stop_words="english", max_features=1000)
    X = vec.fit_transform([text])
    feats = np.array(vec.get_feature_names_out())
    scores = X.toarray().flatten()
    picks = feats[np.argsort(scores)[::-1][:top_n]]
    return [w.title() for w in picks if w.lower() not in BLACKLIST]

# ————— Helper: insert a bullet after a paragraph —————
def insert_bullet_after(heading_para: Paragraph, text: str, template_para: Paragraph = None):
    # Create a new paragraph XML block
    p = OxmlElement("w:p")
    heading_para._p.addnext(p)
    new_para = Paragraph(p, heading_para._parent)

    # Apply style from existing bullet (template)
    if template_para:
        # Copy paragraph style (like bullet indentation)
        new_para.style = template_para.style

        # Create new run and copy font settings
        new_run = new_para.add_run(text)
        if template_para.runs:
            ref_run = template_para.runs[0]
            new_run.font.name = ref_run.font.name
            new_run.font.size = ref_run.font.size
            new_run.bold = ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
    else:
        # Fallback to default list style
        new_para.style = "List Bullet"
        new_para.add_run(text)

    return new_para



# ————— 1) Resume Customizer View —————
def customize_resume_view(request):
    if request.method == "POST":
        form = ResumeForm(request.POST, request.FILES)
        if not form.is_valid():
            return render(request, "index.html", {"form": form})

        resume_file = request.FILES["resume_file"]
        job_text    = form.cleaned_data["job_desc_text"]

        print(job_text)

        # Try GPT extraction
        prompt = (
            "You are a resume assistant. From the following job description, extract a flat list of relevant skills. "
            "This should include both technical skills (e.g. Python, AWS, Docker) and soft/role-related skills "
            "(e.g. leadership, communication, problem-solving). Only include valid skills. "
            "❌ Do NOT include company names, job titles, or vague phrases like 'looking for' or 'expertise'.\n\n"
            "Return the result as a strict JSON array of strings like:\n"
            '["Python", "AWS", "Leadership", "Communication"]\n\n'
            "Job Description:\n" + job_text
        )

        try:
            resp = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": (
                        "Extract ONLY real skills from a job description — both technical (tools, tech, platforms) and soft skills. "
                        "Return a single JSON array of strings. No explanations. No names, titles, or vague terms like 'expertise' or 'looking for'."
                    )},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.0,
            )

            skills = json.loads(resp.choices[0].message.content.strip())
            print("Skills:", skills)
            # Clean and deduplicate
            skills = list({s.strip().title() for s in skills if isinstance(s, str) and s.strip()})
        except (APIConnectionError, KeyError, json.JSONDecodeError, AttributeError):
            skills = extract_keywords_tfidf(job_text, top_n=10)

        # Clean, title-case, dedupe
        clean = []
        for s in skills:
            s2 = s.strip().title()
            if s2 and s2.lower() not in BLACKLIST and s2 not in clean:
                clean.append(s2)

        # Load or convert resume → Document
        if resume_file.name.lower().endswith(".pdf"):
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            tmp.write(resume_file.read()); tmp.close()
            buf = io.BytesIO()
            Converter(tmp.name).convert(buf); buf.seek(0)
            doc = Document(buf)
            try: os.unlink(tmp.name)
            except: pass
        else:
            doc = Document(resume_file)

        # Find existing "Skills" heading only — do not create a new one
        # 3) Find or add heading
        heading = None
        for p in doc.paragraphs:
            if p.text.strip().lower() in ("skills","technical skills"):
                heading = p; break
        if not heading:
            heading = doc.add_heading("Skills", level=1)

        # Remove bullets under Skills until next heading
        # Get first bullet under heading to copy formatting
        template_para = None
        sib = heading._p.getnext()
        while sib is not None:
            styles = sib.xpath('.//w:pStyle')
            val = styles[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "") if styles else ""
            if val.startswith("Heading"):
                break
            if val.startswith("List") or val.lower().startswith("bullet"):
                template_para = Paragraph(sib, heading._parent)
                break
            sib = sib.getnext()


        # Insert new bullets under the existing heading
        for skill in clean:
            insert_bullet_after(heading, skill, template_para=template_para)

        # Return DOCX
        out = io.BytesIO()
        doc.save(out); out.seek(0)
        return HttpResponse(
            out.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="custom_resume.docx"'}
        )

    # GET
    return render(request, "index.html", {"form": ResumeForm()})


# ————— 2) Cover Letter Generator —————
def cover_letter_view(request):
    if request.method == "POST":
        form = CoverLetterForm(request.POST, request.FILES)
        if form.is_valid():
            job_text = form.cleaned_data["job_description"]
            resume_file = form.cleaned_data.get("resume") or request.FILES.get("resume")
            resume_text = ""
            if resume_file:
                if resume_file.name.lower().endswith(".pdf"):
                    reader = PdfReader(resume_file)
                    resume_text = "\n".join(
                        p.extract_text() for p in reader.pages if p.extract_text()
                    )
                else:
                    docx = Document(resume_file)
                    resume_text = "\n".join(p.text for p in docx.paragraphs)

            prompt = (
                f"Write a professional cover letter. Job:\n{job_text}\n\n"
                f"Resume highlights:\n{resume_text}\n\n"
                "Begin with Dear Hiring Manager and close with Sincerely."
            )
            resp = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role":"system","content":"You write cover letters."},
                    {"role":"user","content":prompt}
                ],
                temperature=0.7
            )
            request.session["cover_text"] = resp.choices[0].message.content
            return redirect("cover_letter")
    else:
        form = CoverLetterForm()

    return render(request, "cover_letter_form.html", {"form": form})

# ————— 3) Download Cover as DOCX —————
def download_cover_docx(request):
    cover_text = request.session.get("cover_text", "")
    template   = os.path.join(settings.BASE_DIR, "static/cover_letters/cover_letter_template.docx")
    doc        = Document(template)
    for line in cover_text.splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="CoverLetter.docx"'}
    )

# ————— 4) Download Cover as PDF —————
def download_cover_pdf(request):
    cover_text = request.session.get("cover_text", "")
    html       = render_to_string("cover_letter_pdf.html", {"body": cover_text})
    pdf_buf    = io.BytesIO()
    HTML(string=html).write_pdf(pdf_buf)
    pdf_buf.seek(0)
    return FileResponse(pdf_buf, as_attachment=True, filename="CoverLetter.pdf")
